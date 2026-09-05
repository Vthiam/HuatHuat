from app.services.document_extract import extract_text


def test_extracts_txt(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Plain text content here.")
    assert extract_text(path) == "Plain text content here."


def test_extracts_docx(tmp_path):
    import docx

    path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph of the checklist.")
    document.add_paragraph("Second paragraph mentions section 4.")
    document.save(str(path))

    text = extract_text(path)
    assert "First paragraph of the checklist." in text
    assert "Second paragraph mentions section 4." in text


def test_extracts_pdf(tmp_path):
    import fitz

    path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "This PDF mentions section 4 of the Act.")
    pdf.save(str(path))
    pdf.close()

    text = extract_text(path)
    assert "section 4" in text


def test_unrecognized_extension_returns_empty_string(tmp_path):
    path = tmp_path / "sample.xyz"
    path.write_text("irrelevant")
    assert extract_text(path) == ""
