from types import SimpleNamespace

import pytest

from app.services import document_editor


@pytest.fixture()
def fake_library(tmp_path, monkeypatch):
    law_library = tmp_path / "law_library"
    templates = law_library / "templates"
    templates.mkdir(parents=True)
    monkeypatch.setattr(document_editor, "LAW_LIBRARY_DIR", law_library)
    return {"root": law_library, "templates": templates}


def test_apply_edit_txt_replaces_text(fake_library):
    path = fake_library["templates"] / "checklist.txt"
    path.write_text("Exempt from Parts III to VI where applicable.")

    document = SimpleNamespace(file_path="templates/checklist.txt")
    result = document_editor.apply_edit(document, "Parts III to VI", "Parts 3, 4, 5, 6, 6A and 6B")

    assert result is True
    assert path.read_text() == "Exempt from Parts 3, 4, 5, 6, 6A and 6B where applicable."


def test_apply_edit_txt_returns_false_when_sentence_missing(fake_library):
    path = fake_library["templates"] / "checklist.txt"
    original = "Nothing relevant here."
    path.write_text(original)

    document = SimpleNamespace(file_path="templates/checklist.txt")
    result = document_editor.apply_edit(document, "Parts III to VI", "Parts 3-6")

    assert result is False
    assert path.read_text() == original  # untouched


def test_apply_edit_docx_replaces_text(fake_library):
    import docx

    path = fake_library["templates"] / "checklist.docx"
    d = docx.Document()
    d.add_paragraph("Intro paragraph, unrelated.")
    d.add_paragraph("Exempt from Parts III to VI where applicable.")
    d.save(str(path))

    document = SimpleNamespace(file_path="templates/checklist.docx")
    result = document_editor.apply_edit(document, "Parts III to VI", "Parts 3, 4, 5, 6, 6A and 6B")

    assert result is True
    reopened = docx.Document(str(path))
    texts = [p.text for p in reopened.paragraphs]
    assert "Exempt from Parts 3, 4, 5, 6, 6A and 6B where applicable." in texts
    assert "Intro paragraph, unrelated." in texts  # other paragraphs untouched


def test_apply_edit_returns_false_for_pdf(fake_library):
    document = SimpleNamespace(file_path="templates/checklist.pdf")
    assert document_editor.apply_edit(document, "anything", "anything else") is False


def test_apply_edit_returns_false_for_missing_file(fake_library):
    document = SimpleNamespace(file_path="templates/does_not_exist.txt")
    assert document_editor.apply_edit(document, "anything", "anything else") is False
