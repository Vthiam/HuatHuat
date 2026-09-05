import hashlib
from types import SimpleNamespace

import pytest

from app.services import docx_commenter


def _sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


@pytest.fixture()
def fake_library(tmp_path, monkeypatch):
    law_library = tmp_path / "law_library"
    templates = law_library / "templates"
    reports = law_library / "reports"
    templates.mkdir(parents=True)
    reports.mkdir(parents=True)
    monkeypatch.setattr(docx_commenter, "LAW_LIBRARY_DIR", law_library)
    monkeypatch.setattr(docx_commenter, "REPORTS_DIR", reports)
    return {"root": law_library, "templates": templates, "reports": reports}


def _make_docx(path, paragraphs):
    import docx

    d = docx.Document()
    for text in paragraphs:
        d.add_paragraph(text)
    d.save(str(path))


def test_adds_real_word_comment_and_leaves_original_untouched(fake_library):
    docx_path = fake_library["templates"] / "checklist.docx"
    _make_docx(docx_path, ["Intro.", "Exempt from Parts III to VI where applicable."])
    original_hash = _sha256(docx_path)

    document = SimpleNamespace(file_path="templates/checklist.docx")
    flag = SimpleNamespace(recommendation_text="This no longer matches the amended section 4 wording.")

    output_path = docx_commenter.add_comment_to_docx(document, flag, "Parts III to VI")

    assert output_path is not None
    assert output_path.exists()
    assert output_path.parent.name == "flagged"
    assert _sha256(docx_path) == original_hash  # original never modified

    import docx

    reopened = docx.Document(str(output_path))
    comments = list(reopened.comments)
    assert len(comments) == 1
    assert "amended section 4" in comments[0].text
    assert comments[0].author == "HuatHuat AI"


def test_returns_none_for_non_docx(fake_library):
    document = SimpleNamespace(file_path="templates/checklist.pdf")
    flag = SimpleNamespace(recommendation_text="note")
    assert docx_commenter.add_comment_to_docx(document, flag, "anything") is None


def test_returns_none_when_sentence_not_found(fake_library):
    docx_path = fake_library["templates"] / "checklist.docx"
    _make_docx(docx_path, ["Completely unrelated content."])

    document = SimpleNamespace(file_path="templates/checklist.docx")
    flag = SimpleNamespace(recommendation_text="note")
    result = docx_commenter.add_comment_to_docx(document, flag, "This phrase is not in the document")
    assert result is None


def test_returns_none_when_source_missing(fake_library):
    document = SimpleNamespace(file_path="templates/does_not_exist.docx")
    flag = SimpleNamespace(recommendation_text="note")
    assert docx_commenter.add_comment_to_docx(document, flag, "anything") is None
