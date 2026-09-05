"""Applies an AI-suggested edit to a document's real content -- the one
place in this whole project that rewrites a document's own text, and it
only ever runs when a human explicitly clicks Accept on a Flag (see
routers/flags.py / cli.py's review command). It is never called
automatically.

Scope, per explicit decision: .docx and .txt only. Both support reliable,
clean text replacement. PDFs are excluded -- a PDF's text lives in a
content stream that resists reliable in-place editing without risking
garbled output, so a PDF Flag keeps the existing highlight-only behavior
(pdf_highlighter.py) and always requires a human to edit it themselves
outside this app.
"""
from pathlib import Path

from ..config import LAW_LIBRARY_DIR


def apply_edit(document, original_sentence: str, suggested_replacement: str) -> bool:
    """Returns True if the edit was actually applied. Returns False (does
    nothing) if the file type isn't supported, the file is missing, or
    original_sentence can no longer be found verbatim (e.g. the document
    changed since the Flag was created) -- callers should treat False as
    "nothing to apply", not an error."""
    suffix = Path(document.file_path).suffix.lower()
    full_path = LAW_LIBRARY_DIR / document.file_path
    if not full_path.exists():
        return False

    if suffix == ".txt":
        return _apply_txt(full_path, original_sentence, suggested_replacement)
    if suffix == ".docx":
        return _apply_docx(full_path, original_sentence, suggested_replacement)
    return False


def _apply_txt(path: Path, original_sentence: str, suggested_replacement: str) -> bool:
    text = path.read_text(encoding="utf-8", errors="ignore")
    if original_sentence not in text:
        return False
    path.write_text(text.replace(original_sentence, suggested_replacement, 1), encoding="utf-8")
    return True


def _apply_docx(path: Path, original_sentence: str, suggested_replacement: str) -> bool:
    import docx

    document = docx.Document(str(path))
    for paragraph in document.paragraphs:
        if original_sentence not in paragraph.text:
            continue
        new_text = paragraph.text.replace(original_sentence, suggested_replacement, 1)
        # Simplification, disclosed: rebuilds the paragraph as a single run
        # with the replacement text, rather than splicing individual runs.
        # A text-level swap either way risks losing run-level formatting
        # (bold/italic mid-sentence) for the edited paragraph specifically;
        # every other paragraph in the document is untouched.
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        document.save(str(path))
        return True
    return False
